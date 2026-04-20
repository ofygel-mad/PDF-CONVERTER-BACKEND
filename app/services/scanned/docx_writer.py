"""
Generate a structured Word (.docx) document from a ScannedDocument.

Output structure:
  Section 1 — Metadata (filename, pages, avg confidence, rotations)
  Section 2 — Per-page tables (Word table, header row styled, [?] cells footnoted)
  Section 3 — Appendix (warnings, low-quality pages, struck-through cells)
"""
from __future__ import annotations

import logging
from io import BytesIO

from app.services.scanned.types import OCRCell, ScannedDocument, ScannedTable

log = logging.getLogger(__name__)


def write_docx(document: ScannedDocument) -> bytes:
    """Serialize ScannedDocument to .docx bytes. Returns empty bytes if python-docx unavailable."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError:
        log.warning("docx_writer: python-docx not installed")
        return b""

    doc = Document()

    # ── Metadata section ──
    doc.add_heading(document.source_filename, level=1)
    meta_lines = [
        f"Источник: {document.source_filename}",
        f"Страниц: {len(document.pages)}",
        f"Средняя уверенность OCR: {document.avg_confidence * 100:.1f}%",
    ]
    rotations = [f"{p.rotation_angle:.1f}°" for p in document.pages if abs(p.rotation_angle) > 0.1]
    if rotations:
        meta_lines.append(f"Автоповорот страниц: {', '.join(rotations)}")
    for line in meta_lines:
        doc.add_paragraph(line)

    # ── Per-page tables ──
    appendix_notes: list[str] = []

    for page in document.pages:
        doc.add_heading(f"Страница {page.page_index + 1}", level=2)

        if page.warnings:
            p = doc.add_paragraph(f"⚠ {', '.join(page.warnings)}")
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

        if not page.tables:
            doc.add_paragraph("Таблицы не обнаружены")
            continue

        for table_idx, scanned_table in enumerate(page.tables):
            _add_word_table(doc, scanned_table, table_idx, appendix_notes)

    # ── Appendix ──
    if appendix_notes or document.warnings:
        doc.add_heading("Приложение: предупреждения", level=2)
        for note in document.warnings + appendix_notes:
            doc.add_paragraph(f"• {note}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_word_table(
    doc,
    scanned_table: ScannedTable,
    table_idx: int,
    appendix_notes: list[str],
) -> None:
    from docx.shared import RGBColor, Pt  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    cells = scanned_table.cells
    if not cells:
        return

    rows = max(c.row for c in cells) + 1
    cols = max(c.col for c in cells) + 1

    # Build row→col→text map
    cell_map: dict[tuple[int, int], OCRCell] = {(c.row, c.col): c for c in cells}

    word_table = doc.add_table(rows=rows, cols=cols)
    try:
        word_table.style = "Table Grid"
    except Exception:
        pass

    for ri in range(rows):
        for ci in range(cols):
            cell = word_table.cell(ri, ci)
            ocr_cell = cell_map.get((ri, ci))
            text = ocr_cell.text if ocr_cell else ""
            cell.text = text

            # Style header row
            if ri == scanned_table.header_row_index:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                _set_cell_bg(cell, "1F4E79")

            # Flag low confidence
            if ocr_cell and ocr_cell.struck_through:
                appendix_notes.append(
                    f"Стр. {scanned_table.page_index + 1}, таб. {table_idx + 1}, "
                    f"ячейка [{ri},{ci}]: зачёркнутый текст: «{text[:40]}»"
                )

    doc.add_paragraph("")  # spacer


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background color via OOXML."""
    try:
        from docx.oxml.ns import qn  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    except Exception:
        pass
